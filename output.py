# given a song, output to a word doc or print to terminal
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import os

fonts = dict({
    "standard":{"size": 12, "bold": True, "name": "Arial"},
    "church_lyric":{"size": 12, "bold": False, "name": "Calibri"},
    "church_chord":{"size": 12, "bold": True, "name": "Calibri"},
    "church_title":{"size": 16, "bold": True, "name": "Calibri"}
})

def TerminalOut(song):
    for structure in song:
        print(structure[0])
        for line in structure[1]:
            print(line[1])

class Docx:
    def __init__(self):
        self.doc = docx.Document()
        self.par = self.doc.add_paragraph()
        self.styles = self.doc.styles
        for font in fonts:
            self.addStyle(font, fname = fonts[font]['name'], fsize = fonts[font]['size'], bold = fonts[font]['bold'])

    def addStyle(self, name, fname=False, fsize=False, bold=False, italic=False, underline=False ):
        style = self.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        if fname:
            font.name = fname
        if fsize:
            font.size = Pt(fsize)
        font.bold = bold
        font.italic = italic
        font.underline = underline
        #print(f"'{style.name}' style added")

    def newPar(self):
        self.par = self.doc.add_paragraph()

    def setParStyle(self, name):
        self.newPar()
        self.par.style = self.styles[name]

    def build(self, text, docName=False):
        self.par.add_run(text)
        self.save(docName + ".docx")

    def writeLine(self, text, style=None, end=False):
        text += "\n"
        if style == None:
            self.par.add_run(text)
        elif style == "bold":
            self.par.add_run(text).bold = True
        elif style == "italic":
            self.par.add_run(text).italic = True
        else:
            self.setParStyle(style)
            self.par.add_run(text)
        if end:
            self.save()

    def save(self, docName = False):
        if docName == False:
            docName = input("Enter document name: ") + ".docx"
        else:
            if ".docx" not in docName:
                docName += ".docx"
        while os.path.exists(docName):
            inp = input("Document '" + docName.strip(".docx") + "' already exists, enter a different document name or press enter to overwrite: ")
            if inp == "":
                try:
                    os.remove(docName)
                except:
                    print(f"Document in use, unable to overwrite. Close program currently using '{docName}'.\n")
            else:
                docName = inp + ".docx"
        self.doc.save(docName)
    

if __name__ == '__main__': 
    ins = Docx()
    a = ins.styles['standard']
    ins.setParStyle('standard')
    ins.writeLine("Hello world")
    ins.save()
    print()